from fastapi import FastAPI, HTTPException, BackgroundTasks
from pydantic import BaseModel, Field
from typing import Dict, Any, Optional, List
import logging
import json
import os
from datetime import datetime
from io import BytesIO
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import uvicorn

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(title="Word Document Generator API", version="1.0.0")
os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = r"D:/code/GoogleADK/CloudAssist/add_agent/add_dev/cloud-practice-dev-2-edda3798bdf3.json"

# Pydantic models
class DocumentRequest(BaseModel):
    mal_id: str = Field(..., description="MAL identifier")
    app_name: str = Field(..., description="Application name")
    storage_type: str = Field(default="local", description="Storage type: 'local' or 'gcs'")
    bucket_name: Optional[str] = Field(None, description="GCS bucket name (required if storage_type='gcs')")
    folder_prefix: str = Field(default="documents", description="Folder prefix")
    output_directory: str = Field(default="./generated_documents", description="Local output directory")
    document_data: Dict[str, Any] = Field(..., description="Document content data")


class DocumentResponse(BaseModel):
    status: str
    mal_id: str
    file_path: Optional[str] = None
    gcs_path: Optional[str] = None
    file_size_mb: Optional[float] = None
    processing_time: float
    message: str
    error_details: Optional[str] = None


class FlexibleWordGenerator:
    """
    Word document generator that can save locally or upload to GCS
    """

    def __init__(self, mal_id: str = "default", debug: bool = False):
        self.mal_id = mal_id
        self.debug = debug
        self.logger = self._setup_logger()

    def _setup_logger(self) -> logging.Logger:
        """Setup logger for this generator"""
        logger_name = f"FlexibleWordGen_{self.mal_id}"
        logger = logging.getLogger(logger_name)

        if not logger.handlers:
            handler = logging.StreamHandler()
            formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
            handler.setFormatter(formatter)
            logger.addHandler(handler)

        logger.setLevel(logging.DEBUG if self.debug else logging.INFO)
        return logger

    def _setup_document_styles(self, doc: Document):
        """Setup document styles with professional formatting"""
        try:
            # Set default font
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Cambria'
            font.size = Pt(12)

            # Create heading styles
            heading_1 = doc.styles['Heading 1']
            heading_1.font.size = Pt(16)
            heading_1.font.name = 'Cambria'
            heading_1.font.bold = True

            heading_2 = doc.styles['Heading 2']
            heading_2.font.size = Pt(14)
            heading_2.font.name = 'Cambria'
            heading_2.font.bold = True

            self.logger.debug("Document styles setup completed")
        except Exception as e:
            self.logger.error(f"Error setting up styles: {e}")

    def _create_title_page(self, doc: Document, app_name: str, mal_name: str):
        """Create professional title page"""
        try:
            # Add logo
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run('LUMEN')
            run.font.name = 'Arial Black'
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 100, 200)

            # Add main content
            table = doc.add_table(rows=1, cols=1)
            cell = table.cell(0, 0)

            # Add DRAFT
            for _ in range(4):
                cell.add_paragraph()

            p = cell.add_paragraph('DRAFT')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(40)
            p.runs[0].font.bold = True
            p.runs[0].font.color.rgb = RGBColor(255, 0, 0)

            # Add title
            p = cell.add_paragraph('Application Design Document')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(20)
            p.runs[0].font.bold = True

            # Add app details
            p = cell.add_paragraph(f'Application Name: {app_name}')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(20)
            p.runs[0].font.bold = True

            p = cell.add_paragraph(f'MAL Name: {mal_name}')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(20)
            p.runs[0].font.bold = True

            # Add document details
            for _ in range(8):
                cell.add_paragraph()

            current_date = datetime.now().strftime('%Y-%m-%d')
            details = [
                'Document Reference: ADD-001',
                'Version: 1.4',
                'Document Status: Draft submitted for review',
                'Author: System Generated',
                f'Last Revision Date: {current_date}',
                f'Date Created: {current_date}'
            ]

            for detail in details:
                p = cell.add_paragraph(detail)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.runs[0].font.bold = True
                p.runs[0].font.size = Pt(11)

            doc.add_page_break()
            self.logger.debug("Title page created")
        except Exception as e:
            self.logger.error(f"Error creating title page: {e}")

    def _add_table_of_contents(self, doc: Document):
        """Add table of contents"""
        try:
            doc.add_heading('Table of Contents', level=1)

            toc_items = [
                '1. Application Overview',
                '2. Technical Architecture',
                '3. Development and Operations',
                '4. Risks, Assumptions, and Decisions',
                '5. Appendix'
            ]

            for item in toc_items:
                doc.add_paragraph(item)

            doc.add_page_break()
            self.logger.debug("Table of contents added")
        except Exception as e:
            self.logger.error(f"Error creating table of contents: {e}")

    def _create_simple_table(self, doc: Document, title: str, headers: List[str], data: List[Dict]):
        """Create a simple table with data"""
        try:
            if not data:
                return

            doc.add_heading(title, level=3)
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'

            # Add headers
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            # Add data rows
            for row_data in data:
                row_cells = table.add_row().cells
                for i, header in enumerate(headers):
                    value = row_data.get(header, row_data.get(header.lower(), ''))
                    row_cells[i].text = str(value) if value else ""

            self.logger.debug(f"Table created: {title}")
        except Exception as e:
            self.logger.error(f"Error creating table {title}: {e}")

    def _process_document_content(self, doc: Document, app_name: str, document_data: Dict):
        """Process document data into Word content"""
        try:
            # Application Overview
            doc.add_heading('1. Application Overview', level=1)
            overview = document_data.get('overview',
                                         f'This document describes the architecture design for the {app_name} application.')
            doc.add_paragraph(overview)

            # Technical Architecture
            if 'technical_architecture' in document_data:
                doc.add_heading('2. Technical Architecture', level=1)
                tech_data = document_data['technical_architecture']

                # Technology Stack
                tech_stack = tech_data.get('technology_stack', 'Not specified')
                doc.add_heading('2.1 Technology Stack', level=2)
                doc.add_paragraph(f'Technology: {tech_stack}')

                # Components
                if 'components' in tech_data:
                    doc.add_heading('2.2 System Components', level=2)
                    for component in tech_data['components']:
                        doc.add_paragraph(f"• {component}")

            # Custom sections
            if 'sections' in document_data:
                for section in document_data['sections']:
                    title = section.get('title', 'Custom Section')
                    level = section.get('level', 2)
                    content = section.get('content', '')

                    doc.add_heading(title, level=level)
                    doc.add_paragraph(content)

            # Tables
            if 'tables' in document_data:
                for table_data in document_data['tables']:
                    title = table_data.get('title', 'Data Table')
                    headers = table_data.get('headers', [])
                    rows = table_data.get('data', [])
                    if headers and rows:
                        self._create_simple_table(doc, title, headers, rows)

            self.logger.debug("Document content processed successfully")
        except Exception as e:
            self.logger.error(f"Error processing document content: {e}")

    def _save_locally(self, doc: Document, app_name: str, output_directory: str, folder_prefix: str) -> Dict[str, Any]:
        """Save document to local filesystem"""
        try:
            # Create output directory
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            full_output_path = Path(output_directory) / folder_prefix / self.mal_id / timestamp
            full_output_path.mkdir(parents=True, exist_ok=True)

            # Create filename
            filename = f"{app_name.replace(' ', '_')}_Document_{timestamp}.docx"
            file_path = full_output_path / filename

            # Save document
            doc.save(str(file_path))

            # Get file size
            file_size = file_path.stat().st_size

            self.logger.info(f"Document saved locally: {file_path}")

            return {
                'status': 'success',
                'file_path': str(file_path),
                'file_size_bytes': file_size,
                'file_size_mb': round(file_size / (1024 * 1024), 2),
                'filename': filename,
                'message': 'Document saved successfully to local filesystem'
            }

        except Exception as e:
            self.logger.error(f"Failed to save document locally: {e}")
            raise e

    def _save_to_gcs(self, doc: Document, app_name: str, bucket_name: str, folder_prefix: str) -> Dict[str, Any]:
        """Save document to GCS"""
        try:
            from google.cloud import storage
        except ImportError:
            raise ImportError("google-cloud-storage not installed. Run: pip install google-cloud-storage")

        # Create filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{app_name.replace(' ', '_')}_Document_{timestamp}.docx"

        # Save document to BytesIO (in memory)
        doc_buffer = BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)

        # Get document size
        doc_size = doc_buffer.getbuffer().nbytes

        try:
            client = storage.Client()
            bucket = client.bucket(bucket_name)

            # Create blob path
            blob_name = f"{folder_prefix}/{self.mal_id}/{timestamp}/{filename}"
            blob = bucket.blob(blob_name)

            # Upload from memory buffer
            blob.upload_from_file(
                doc_buffer,
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

            gcs_path = f"gs://{bucket_name}/{blob_name}"

            self.logger.info(f"Successfully uploaded: {gcs_path}")

            return {
                'status': 'success',
                'gcs_path': gcs_path,
                'file_size_bytes': doc_size,
                'file_size_mb': round(doc_size / (1024 * 1024), 2),
                'filename': filename,
                'message': 'Document uploaded successfully to GCS'
            }

        except Exception as upload_error:
            self.logger.error(f"Failed to upload {filename}: {upload_error}")
            raise upload_error

        finally:
            doc_buffer.close()

    def generate_document(self, app_name: str, document_data: Dict,
                          storage_type: str = "local", **kwargs) -> Dict[str, Any]:
        """Generate Word document and save to specified storage"""
        start_time = datetime.now()

        try:
            # Validate inputs
            if not document_data:
                raise ValueError("document_data is required")

            self.logger.info(f"Generating document for: {app_name}")

            # Create document in memory
            doc = Document()
            self._setup_document_styles(doc)

            # Build document content
            self._create_title_page(doc, app_name, self.mal_id)
            self._add_table_of_contents(doc)
            self._process_document_content(doc, app_name, document_data)

            # Save based on storage type
            if storage_type.lower() == "gcs":
                bucket_name = kwargs.get('bucket_name')
                if not bucket_name:
                    raise ValueError("bucket_name is required for GCS storage")

                folder_prefix = kwargs.get('folder_prefix', 'documents')
                result = self._save_to_gcs(doc, app_name, bucket_name, folder_prefix)

            else:  # local storage
                output_directory = kwargs.get('output_directory', './generated_documents')
                folder_prefix = kwargs.get('folder_prefix', 'documents')
                result = self._save_locally(doc, app_name, output_directory, folder_prefix)

            # Add processing time
            processing_time = (datetime.now() - start_time).total_seconds()
            result['processing_time'] = processing_time

            return result

        except Exception as e:
            processing_time = (datetime.now() - start_time).total_seconds()
            self.logger.error(f"Document generation failed: {e}")

            return {
                'status': 'error',
                'processing_time': processing_time,
                'error': str(e),
                'message': f'Document generation failed: {str(e)}'
            }


# API Endpoints
@app.post("/generate-document", response_model=DocumentResponse)
async def generate_document(request: DocumentRequest):
    """
    Generate a Word document and save it locally or upload to GCS
    """
    try:
        logger.info(f"Received document generation request for MAL ID: {request.mal_id}")

        # Create generator
        generator = FlexibleWordGenerator(mal_id=request.mal_id, debug=False)

        # Prepare kwargs based on storage type
        kwargs = {
            'folder_prefix': request.folder_prefix
        }

        if request.storage_type.lower() == "gcs":
            if not request.bucket_name:
                raise HTTPException(status_code=400, detail="bucket_name is required for GCS storage")
            kwargs['bucket_name'] = request.bucket_name
        else:
            kwargs['output_directory'] = request.output_directory

        # Generate document
        result = generator.generate_document(
            app_name=request.app_name,
            document_data=request.document_data,
            storage_type=request.storage_type,
            **kwargs
        )

        if result['status'] == 'success':
            return DocumentResponse(
                status="success",
                mal_id=request.mal_id,
                file_path=result.get('file_path'),
                gcs_path=result.get('gcs_path'),
                file_size_mb=result['file_size_mb'],
                processing_time=result['processing_time'],
                message=result['message']
            )
        else:
            raise HTTPException(
                status_code=500,
                detail=f"Document generation failed: {result.get('message', 'Unknown error')}"
            )

    except Exception as e:
        logger.error(f"API error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {"status": "healthy", "timestamp": datetime.now().isoformat()}


@app.get("/")
async def root():
    """Root endpoint with API information"""
    return {
        "message": "Flexible Word Document Generator API",
        "version": "1.0.0",
        "storage_options": ["local", "gcs"],
        "endpoints": {
            "generate_document": "/generate-document",
            "health": "/health",
            "docs": "/docs"
        }
    }


if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)
