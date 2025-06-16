from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from datetime import datetime
import html.parser
import re
import os
import requests
from io import BytesIO
from PIL import Image
import base64


class WordGenerator:
    def __init__(self):
        self.doc = Document()
        self.setup_styles()

    def setup_styles(self):
        """Setup document styles"""
        # Set default font
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Cambria'
        font.size = Pt(12)

        # Create heading styles
        heading_1 = self.doc.styles['Heading 1']
        heading_1.font.size = Pt(16)
        heading_1.font.name = 'Cambria'

        heading_2 = self.doc.styles['Heading 2']
        heading_2.font.size = Pt(14)
        heading_2.font.name = 'Cambria'

    def add_border_to_paragraph(self, paragraph):
        """Add border to paragraph"""
        p = paragraph._element
        pBdr = OxmlElement('w:pBdr')

        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '25')
        top.set(qn('w:space'), '1')
        top.set(qn('w:color'), '0000FF')

        pBdr.append(top)
        p.get_or_add_pPr().append(pBdr)

    def convert_permissions_to_paragraphs(self, permissions_string, table_cell):
        """Convert comma-separated permissions to paragraphs in a table cell"""
        if not permissions_string:
            return

        permissions = [p.strip() for p in permissions_string.split(',')]
        for i, permission in enumerate(permissions):
            p = table_cell.add_paragraph()
            if i < len(permissions) - 1:
                p.add_run(permission + ',')
            else:
                p.add_run(permission)

    def convert_html_to_paragraphs(self, html_string):
        """Convert HTML string to Word paragraphs with proper formatting"""
        if not html_string:
            return []

        from bs4 import BeautifulSoup

        soup = BeautifulSoup(html_string, 'html.parser')

        def process_element(element, paragraph=None):
            if paragraph is None:
                paragraph = self.doc.add_paragraph()

            if element.name == 'b' or element.name == 'strong':
                run = paragraph.add_run(element.get_text())
                run.bold = True
            elif element.name == 'i' or element.name == 'em':
                run = paragraph.add_run(element.get_text())
                run.italic = True
            elif element.name == 'u':
                run = paragraph.add_run(element.get_text())
                run.underline = True
            elif element.name == 'br':
                paragraph.add_run('\n')
            elif element.name == 'p':
                new_p = self.doc.add_paragraph()
                for child in element.children:
                    if child.name:
                        process_element(child, new_p)
                    else:
                        # Text node
                        process_text(str(child), new_p)
            elif element.name in ['ul', 'ol']:
                for li in element.find_all('li'):
                    p = self.doc.add_paragraph('• ' + li.get_text())
                    p.paragraph_format.left_indent = Inches(0.5)
            else:
                # Default: just get text
                process_text(element.get_text(), paragraph)

        def process_text(text, paragraph):
            # Handle special $...$ formatting
            parts = re.split(r'(\$[^$]+\$)', text)
            for part in parts:
                if part.startswith('$') and part.endswith('$'):
                    content = part[1:-1]
                    run = paragraph.add_run(content)
                    run.font.color.rgb = RGBColor(255, 127, 0)  # Orange
                    run.font.size = Pt(12)
                else:
                    run = paragraph.add_run(part)
                    run.font.size = Pt(12)

        # Process root elements
        for element in soup.children:
            if element.name:
                process_element(element)
            else:
                # Direct text
                p = self.doc.add_paragraph()
                process_text(str(element), p)

        return []

    def process_table_cell_content(self, html_content, cell):
        """Process HTML content for table cells"""
        if not html_content:
            return

        # Remove HTML tags
        text = re.sub('<[^<]+?>', '', html_content)

        # Split by dashes
        lines = [line.strip() for line in text.split('-') if line.strip()]

        for i, line in enumerate(lines):
            p = cell.add_paragraph()

            # Handle $...$ formatting
            parts = re.split(r'(\$[^$]+\$)', line)
            for part in parts:
                if part.startswith('$') and part.endswith('$'):
                    content = part[1:-1]
                    run = p.add_run(content)
                    run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
                else:
                    run = p.add_run(part)

            if i > 0:
                p.paragraph_format.space_before = Pt(6)

    def create_table(self, headers, data, key_mapping):
        """Create a table with headers and data"""
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'

        # Add headers
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header

        # Add data rows
        for row_data in data:
            row_cells = table.add_row().cells
            for i, key in enumerate(key_mapping):
                if isinstance(row_data, dict):
                    value = row_data.get(key, '')
                    if key == 'permission' and value:
                        self.convert_permissions_to_paragraphs(value, row_cells[i])
                    else:
                        row_cells[i].text = str(value)

        return table

    def add_image_from_file(self, image_path, width=None, height=None):
        """Add image from file path"""
        if width and height:
            self.doc.add_picture(image_path, width=Inches(width), height=Inches(height))
        else:
            self.doc.add_picture(image_path)

    def add_image_from_url(self, image_url, width=None, height=None):
        """Add image from URL"""
        response = requests.get(image_url)
        image = Image.open(BytesIO(response.content))

        # Save temporarily
        temp_path = 'temp_image.png'
        image.save(temp_path)

        # Add to document
        self.add_image_from_file(temp_path, width, height)

        # Clean up
        os.remove(temp_path)

    def add_lumen_logo(self, paragraph):
        """Add LUMEN text as logo"""
        run = paragraph.add_run('LUMEN')
        run.font.name = 'Arial Black'
        run.font.size = Pt(16)
        run.font.bold = True
        # Optional: make it blue like in many corporate logos
        # run.font.color.rgb = RGBColor(0, 0, 255)

    def create_title_page(self, data):
        """Create the title page"""
        # Add logo at the top right
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self.add_lumen_logo(p)

        # Add the main table for title page content
        table = self.doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'

        # Get the cell
        cell = table.cell(0, 0)

        # Set cell background color to light gray
        from docx.oxml import parse_xml
        shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
        cell._element.get_or_add_tcPr().append(shading_elm)

        # Add vertical spacing before DRAFT
        for _ in range(4):
            cell.add_paragraph()

        # Add DRAFT
        p = cell.add_paragraph('DRAFT')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(40)
        p.runs[0].font.bold = True
        p.runs[0].font.name = 'Cambria'

        # Add spacing
        cell.add_paragraph()

        # Add "Application Design Document"
        p = cell.add_paragraph('Application Design Document')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(20)
        p.runs[0].font.bold = True
        p.runs[0].font.name = 'Cambria'

        # Add Mass Market Unit
        p = cell.add_paragraph(f'Mass Market Unit : {data.get("mmUnit", "")}')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(20)
        p.runs[0].font.bold = True
        p.runs[0].font.name = 'Cambria'

        # Add Application Name
        p = cell.add_paragraph(f'Application Name: {data.get("apName", "")}')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(20)
        p.runs[0].font.bold = True
        p.runs[0].font.name = 'Cambria'

        # Add MAL Name
        p = cell.add_paragraph(f'MAL Name: {data.get("mName", "")}')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(20)
        p.runs[0].font.bold = True
        p.runs[0].font.name = 'Cambria'

        # Add significant vertical spacing
        for _ in range(10):
            cell.add_paragraph()

        # Document details section (left-aligned)
        details = [
            'Document Reference:',
            'Version: 1.4',
            'Document Status: Draft submitted for review',
            f'Author: {data.get("authorName", "")}',
            f'Last Revision Date: {data.get("formattedDate", "")}',
            f'Date Created: {data.get("formattedDate", "")}',
            'Date Printed:'
        ]

        for detail in details:
            p = cell.add_paragraph(detail)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.runs[0].font.bold = True
            p.runs[0].font.size = Pt(11)
            p.runs[0].font.name = 'Cambria'

        # Set the table width to full page width
        table.autofit = False
        table.width = self.doc.sections[0].page_width - self.doc.sections[0].left_margin - self.doc.sections[
            0].right_margin

        # Page break
        self.doc.add_page_break()

    def add_table_of_contents(self):
        """Add table of contents"""
        self.doc.add_heading('Table of Contents', level=1)
        # Note: python-docx doesn't support automatic TOC generation
        # You would need to manually add TOC entries or use a different approach
        self.doc.add_paragraph('1. Application Overview')
        self.doc.add_paragraph('    1.1 Business Needs')
        self.doc.add_paragraph('    1.2 Key Features')
        self.doc.add_paragraph('    1.3 User Stories')
        self.doc.add_paragraph('2. Technical Architecture')
        self.doc.add_paragraph('    2.1 Logical Architecture')
        self.doc.add_paragraph('    2.2 Technology Stack')
        self.doc.add_paragraph('    2.3 Deployment Model')
        self.doc.add_paragraph('    2.4 Project Access')
        self.doc.add_paragraph('    2.5 Key Components')
        self.doc.add_paragraph('3. Development Operations')
        self.doc.add_paragraph('4. Risk, Assumptions and Decisions')
        self.doc.add_paragraph('5. Appendix')
        self.doc.add_page_break()

    def add_headers_footers(self):
        """Add headers and footers to the document"""
        section = self.doc.sections[0]

        # Header
        header = section.header
        header_paragraph = header.paragraphs[0]
        header_paragraph.text = "High Level Architecture Design Document"
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Add border to header
        self.add_border_to_paragraph(header_paragraph)

        # Footer
        footer = section.footer

        # Add border to footer
        footer_paragraph = footer.add_paragraph()
        self.add_border_to_paragraph(footer_paragraph)

        # Footer content
        p1 = footer.add_paragraph("Draft\nInternal Use Only")
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT

        p2 = footer.add_paragraph("Architecture Design Document ©2024 Lumen Technologies. All Rights Reserved.")
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Page numbers (python-docx doesn't support automatic page numbers easily)
        p3 = footer.add_paragraph("Page ")
        p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def generate_document(self, data):
        """Main method to generate the complete document"""
        # Create title page
        self.create_title_page(data)

        # Add table of contents
        self.add_table_of_contents()

        # 1. Application Overview
        self.doc.add_heading('1. Application Overview', level=1)
        self.convert_html_to_paragraphs(data.get('htmlAppOverviewData', ''))

        # 1.1 Business Needs
        self.doc.add_heading('1.1 Business Needs', level=2)
        self.convert_html_to_paragraphs(data.get('htmlBNData', ''))

        # 1.2 Key Features
        self.doc.add_heading('1.2 Key Features', level=2)
        self.convert_html_to_paragraphs(data.get('htmlKFData', ''))

        # 1.3 User Stories
        self.doc.add_heading('1.3 User Stories', level=2)
        self.convert_html_to_paragraphs(data.get('htmlUSData', ''))

        # 2. Technical Architecture
        self.doc.add_heading('2. Technical Architecture', level=1)

        # 2.1 Logical Architecture
        self.doc.add_heading('2.1 Logical Architecture', level=2)
        self.convert_html_to_paragraphs(data.get('htmlLAData', ''))

        # Add images if available
        images = data.get('images', [])
        for image in images:
            if 'preview' in image:
                try:
                    self.add_image_from_url(image['preview'])
                except Exception as e:
                    print(f"Error adding image: {e}")

        # 2.2 Technology Stack
        self.doc.add_heading('2.2 Technology Stack', level=2)
        ts_data = data.get('tsData', '')
        if ts_data:
            # Convert to bullet points
            items = ts_data.split(',')
            for item in items:
                p = self.doc.add_paragraph(f'• {item.strip()}')

        # 2.3 Deployment Model
        self.doc.add_heading('2.3 Deployment Model', level=2)
        self.convert_html_to_paragraphs(data.get('htmlDMData', ''))

        # 2.4 Project Access
        self.doc.add_heading('2.4 Project Access', level=2)
        pa_data = data.get('paData', [])
        if pa_data:
            headers = ['Group', 'Owners', 'Description', 'Environment']
            key_mapping = ['group', 'owners', 'description', 'environment']
            self.create_table(headers, pa_data, key_mapping)

        # 2.5 Key Components
        self.doc.add_heading('2.5 Key Components', level=2)
        kc_data = data.get('htmlKCDataTable', [])
        if kc_data:
            headers = ['Service', 'Dev', 'Prod', 'QA']
            key_mapping = ['service', 'dev', 'prod', 'qa']
            table = self.create_table(headers, kc_data, key_mapping)

            # Process HTML content in cells if needed
            for i, row_data in enumerate(kc_data):
                row = table.rows[i + 1]  # Skip header row
                for j, key in enumerate(key_mapping):
                    cell = row.cells[j]
                    html_content = row_data.get(key, '')
                    if html_content and '$' in html_content:
                        self.process_table_cell_content(html_content, cell)

        # Add special notes if GKE is in technology stack
        if 'GKE' in data.get('tsData', ''):
            p = self.doc.add_paragraph(
                '** Identify current load balancers and balancing strategy (round robin, etc) to distribute traffic to different instances?')
            p.runs[0].font.color.rgb = RGBColor(255, 0, 0)

            p = self.doc.add_paragraph('    ** Kubernetes Ingress level Load Balancer (nginx)')
            p.runs[0].font.color.rgb = RGBColor(255, 0, 0)

            p = self.doc.add_paragraph('    ** Least Connection (Member) Balancing Strategy')
            p.runs[0].font.color.rgb = RGBColor(255, 0, 0)

            p = self.doc.add_paragraph('** Are you relying on the load balancer in K8? Yes')
            p.runs[0].font.color.rgb = RGBColor(255, 0, 0)

        # 2.5.4 Firewall Rules
        self.doc.add_heading('2.5.4 Firewall Rules', level=3)
        fr_data = data.get('frData', [])
        if fr_data:
            headers = ['Network', 'Actions', 'Source IP', 'Destination', 'Port', 'Protocol', 'Direction']
            key_mapping = ['network', 'actions', 'sourceip', 'destination', 'port', 'protocol', 'direction']
            self.create_table(headers, fr_data, key_mapping)

        # 2.5.6 IAM
        self.doc.add_heading('2.5.6 IAM', level=3)
        iam_data = data.get('iamData', [])
        if iam_data:
            headers = ['Service', 'Principal', 'Environment', 'Permission']
            key_mapping = ['service', 'principal', 'environment', 'permission']
            self.create_table(headers, iam_data, key_mapping)

        # Add all other sections as needed...
        # (I'll continue with a sample of the remaining sections)

        # 3. Development Operations
        self.doc.add_heading('3. Development Operations', level=1)

        # 3.3 Operational Support Model
        self.doc.add_heading('3.3 Operational Support Model', level=2)
        osm_data = data.get('osmData', [])
        if osm_data:
            headers = ['Component', 'DEV', 'PROD', 'QA', 'UAT']
            key_mapping = ['component', 'dev', 'prod', 'qa', 'uat']
            self.create_table(headers, osm_data, key_mapping)

        # Add headers and footers
        self.add_headers_footers()

        # Save the document
        filename = f"{data.get('mName', 'Document')} - Architecture Design Document MMCE V1.0.docx"
        self.doc.save(filename)

        # Upload to GCS (if needed)
        self.upload_to_gcs(filename, data.get('fileName', ''))

        return filename

    def upload_to_gcs(self, filename, excel_filename):
        """Upload file to Google Cloud Storage"""
        try:
            from google.cloud import storage

            # Initialize GCS client
            client = storage.Client()

            # Get bucket (you'll need to configure this)
            bucket_name = os.environ.get('GCS_BUCKET_NAME', 'your-bucket-name')
            bucket = client.bucket(bucket_name)

            # Create blob name
            blob_name = f"documents/{excel_filename}/{filename}"
            blob = bucket.blob(blob_name)

            # Upload file
            with open(filename, 'rb') as f:
                blob.upload_from_file(f)

            print(f"File uploaded to GCS: gs://{bucket_name}/{blob_name}")

            # Optional: make the blob publicly accessible
            # blob.make_public()

            return blob.public_url if hasattr(blob, 'public_url') else f"gs://{bucket_name}/{blob_name}"

        except Exception as e:
            print(f"Error uploading to GCS: {e}")
            return None


# Usage example
def main():
    generator = WordGenerator()

    # Sample data
    data = {
        'mmUnit': 'Sample Unit',
        'apName': 'Sample Application',
        'mName': 'Sample MAL',
        'authorName': 'John Doe',
        'formattedDate': datetime.now().strftime('%Y-%m-%d'),
        'htmlAppOverviewData': '<p>This is the application overview</p>',
        'htmlBNData': '<p>Business needs content</p>',
        'tsData': 'Python, Docker, GKE, PostgreSQL',
        'paData': [
            {'group': 'Admin', 'owners': 'John Doe', 'description': 'Admin access', 'environment': 'Prod'},
            {'group': 'Dev', 'owners': 'Jane Smith', 'description': 'Developer access', 'environment': 'Dev'}
        ],
        'frData': [
            {
                'network': '10.0.0.0/24',
                'actions': 'Allow',
                'sourceip': '0.0.0.0',
                'destination': '10.0.0.1',
                'port': '443',
                'protocol': 'TCP',
                'direction': 'Ingress'
            }
        ]
    }

    filename = generator.generate_document(data)
    print(f"Document generated: {filename}")


if __name__ == "__main__":
    main()
