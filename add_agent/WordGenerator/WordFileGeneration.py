from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from datetime import datetime
import html.parser
import re
import os
import requests
from io import BytesIO
from PIL import Image
import base64
import json
import logging
from typing import List, Dict, Optional, Any
from pathlib import Path


class WordGenerator:
    """
    Professional Word document generator for Application Design Documents (ADD)

    This class handles the complete generation of Word documents from ADD data,
    including proper formatting, tables, and GCS upload functionality.
    """

    def __init__(self, debug: bool = False):
        self.doc = Document()
        self.setup_styles()
        self.debug = debug
        self.logger = self._setup_logger()

    def _setup_logger(self) -> logging.Logger:
        """Setup logging for the WordGenerator"""
        logger = logging.getLogger(self.__class__.__name__)
        if not logger.handlers:
            handler = logging.StreamHandler()
            formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
            handler.setFormatter(formatter)
            logger.addHandler(handler)
            logger.setLevel(logging.DEBUG if self.debug else logging.INFO)
        return logger

    def setup_styles(self):
        """Setup document styles with professional formatting"""
        try:
            # Set default font
            style = self.doc.styles['Normal']
            font = style.font
            font.name = 'Cambria'
            font.size = Pt(12)

            # Create heading styles
            heading_1 = self.doc.styles['Heading 1']
            heading_1.font.size = Pt(16)
            heading_1.font.name = 'Cambria'
            heading_1.font.bold = True

            heading_2 = self.doc.styles['Heading 2']
            heading_2.font.size = Pt(14)
            heading_2.font.name = 'Cambria'
            heading_2.font.bold = True

            heading_3 = self.doc.styles['Heading 3']
            heading_3.font.size = Pt(13)
            heading_3.font.name = 'Cambria'
            heading_3.font.bold = True

            heading_4 = self.doc.styles.add_style('Heading 4', WD_STYLE_TYPE.PARAGRAPH)
            heading_4.base_style = self.doc.styles['Normal']
            heading_4.font.size = Pt(12)
            heading_4.font.name = 'Cambria'
            heading_4.font.bold = True

            self.logger.debug("Document styles setup completed")

        except Exception as e:
            self.logger.error(f"Error setting up styles: {e}")

    def add_border_to_paragraph(self, paragraph):
        """Add border to paragraph for headers/footers"""
        try:
            p = paragraph._element
            pBdr = OxmlElement('w:pBdr')

            top = OxmlElement('w:top')
            top.set(qn('w:val'), 'single')
            top.set(qn('w:sz'), '25')
            top.set(qn('w:space'), '1')
            top.set(qn('w:color'), '0000FF')

            pBdr.append(top)
            p.get_or_add_pPr().append(pBdr)

        except Exception as e:
            self.logger.error(f"Error adding border to paragraph: {e}")

    def add_formatted_text_to_cell(self, text: str, cell):
        """
        Add formatted text to a table cell, handling special formatting

        Args:
            text: Text to add to the cell
            cell: Table cell object
        """
        if not text:
            return

        try:
            # Clear the cell
            cell.text = ""

            # Handle comma-separated permissions (long text)
            if ',' in text and len(text) > 50:
                permissions = [p.strip() for p in text.split(',')]
                for i, permission in enumerate(permissions):
                    p = cell.add_paragraph()
                    if i < len(permissions) - 1:
                        p.add_run(permission + ',')
                    else:
                        p.add_run(permission)
                    p.paragraph_format.space_after = Pt(2)
            else:
                # Handle special formatting like $...$
                parts = re.split(r'(\$[^$]+\$)', text)
                p = cell.add_paragraph()
                for part in parts:
                    if part.startswith('$') and part.endswith('$'):
                        content = part[1:-1]
                        run = p.add_run(content)
                        run.font.color.rgb = RGBColor(255, 127, 0)  # Orange color
                        run.font.bold = True
                    else:
                        run = p.add_run(part)

        except Exception as e:
            self.logger.error(f"Error formatting text in cell: {e}")
            # Fallback to plain text
            cell.text = str(text)

    def add_lumen_logo(self, paragraph):
        """Add LUMEN text as company logo"""
        try:
            run = paragraph.add_run('LUMEN')
            run.font.name = 'Arial Black'
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 100, 200)  # Corporate blue

        except Exception as e:
            self.logger.error(f"Error adding logo: {e}")

    def create_title_page(self, app_name: str, mal_name: str):
        """
        Create professional title page

        Args:
            app_name: Application name
            mal_name: MAL name
        """
        try:
            # Add logo at the top right
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            self.add_lumen_logo(p)

            # Add the main table for title page content
            table = self.doc.add_table(rows=1, cols=1)
            table.style = 'Table Grid'

            # Get the cell
            cell = table.cell(0, 0)

            # Set cell background color to light gray
            shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
            cell._element.get_or_add_tcPr().append(shading_elm)

            # Add vertical spacing before DRAFT
            for _ in range(4):
                cell.add_paragraph()

            # Add DRAFT
            p = cell.add_paragraph('DRAFT')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(40)
            p.runs[0].font.bold = True
            p.runs[0].font.name = 'Cambria'
            p.runs[0].font.color.rgb = RGBColor(255, 0, 0)  # Red color for DRAFT

            # Add spacing
            cell.add_paragraph()

            # Add "Application Design Document"
            p = cell.add_paragraph('Application Design Document')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(20)
            p.runs[0].font.bold = True
            p.runs[0].font.name = 'Cambria'

            # Add Mass Market Unit
            p = cell.add_paragraph('Mass Market Unit : Engineering')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(20)
            p.runs[0].font.bold = True
            p.runs[0].font.name = 'Cambria'

            # Add Application Name
            p = cell.add_paragraph(f'Application Name: {app_name}')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(20)
            p.runs[0].font.bold = True
            p.runs[0].font.name = 'Cambria'

            # Add MAL Name
            p = cell.add_paragraph(f'MAL Name: {mal_name}')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(20)
            p.runs[0].font.bold = True
            p.runs[0].font.name = 'Cambria'

            # Add significant vertical spacing
            for _ in range(10):
                cell.add_paragraph()

            # Document details section (left-aligned)
            current_date = datetime.now().strftime('%Y-%m-%d')
            details = [
                'Document Reference: ADD-001',
                'Version: 1.4',
                'Document Status: Draft submitted for review',
                'Author: System Generated',
                f'Last Revision Date: {current_date}',
                f'Date Created: {current_date}',
                f'Date Printed: {current_date}'
            ]

            for detail in details:
                p = cell.add_paragraph(detail)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.runs[0].font.bold = True
                p.runs[0].font.size = Pt(11)
                p.runs[0].font.name = 'Cambria'

            # Set the table width to full page width
            table.autofit = False
            table.width = (self.doc.sections[0].page_width -
                           self.doc.sections[0].left_margin -
                           self.doc.sections[0].right_margin)

            # Page break
            self.doc.add_page_break()
            self.logger.debug(f"Title page created for {app_name}")

        except Exception as e:
            self.logger.error(f"Error creating title page: {e}")

    def add_table_of_contents(self):
        """Add comprehensive table of contents"""
        try:
            self.doc.add_heading('Table of Contents', level=1)

            toc_items = [
                '1. Application Overview',
                '    1.1 Purpose and Scope',
                '    1.2 Business Context',
                '2. Technical Architecture',
                '    2.1 Technology Stack',
                '    2.2 Project Access',
                '    2.3 Key Components',
                '    2.4 Service Accounts',
                '    2.5 Firewall Rules',
                '    2.6 IAM Configuration',
                '3. Development and Operations',
                '    3.1 Development Environment',
                '    3.2 Deployment Process',
                '    3.3 Operational Support Model',
                '    3.4 Service Level Management',
                '    3.5 Governance and Cost Management',
                '4. Risks, Assumptions, and Decisions',
                '    4.1 Design Assumptions',
                '    4.2 Dependencies',
                '    4.3 Design Risks',
                '    4.4 Key Decisions',
                '5. Appendix',
                '    5.1 Abbreviations and Glossary',
                '    5.2 References',
                '    5.3 Resources and Guidance'
            ]

            for item in toc_items:
                p = self.doc.add_paragraph(item)
                # Add some spacing for readability
                if not item.startswith('    '):
                    p.paragraph_format.space_before = Pt(6)

            self.doc.add_page_break()
            self.logger.debug("Table of contents added")

        except Exception as e:
            self.logger.error(f"Error creating table of contents: {e}")

    def create_professional_table(self, headers: List[str], data: List[Dict], title: str = None) -> bool:
        """
        Create a professional table with proper formatting

        Args:
            headers: List of column headers
            data: List of dictionaries containing row data
            title: Optional table title

        Returns:
            bool: Success status
        """
        try:
            if not data:
                self.logger.warning(f"No data provided for table: {title}")
                return False

            if title:
                self.doc.add_paragraph(title, style='Heading 4')

            table = self.doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'

            # Style the table
            table.autofit = True

            # Add headers with formatting
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header
                # Make header text bold
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                # Set header background color
                shading_elm = parse_xml(r'<w:shd {} w:fill="E7E6E6"/>'.format(nsdecls('w')))
                header_cells[i]._element.get_or_add_tcPr().append(shading_elm)

            # Add data rows
            for row_data in data:
                row_cells = table.add_row().cells
                for i, header in enumerate(headers):
                    # Get value using header as key, with fallback options
                    key_variations = [
                        header,
                        header.lower(),
                        header.replace(' ', '_').lower(),
                        header.replace(' ', '').lower()
                    ]

                    value = ""
                    for key in key_variations:
                        if key in row_data:
                            value = row_data[key]
                            break

                    if isinstance(value, (list, dict)):
                        value = str(value)

                    # Apply formatting based on content
                    if value and (',' in str(value)) and len(str(value)) > 50:
                        self.add_formatted_text_to_cell(str(value), row_cells[i])
                    else:
                        row_cells[i].text = str(value) if value else ""

            self.logger.debug(f"Table created with {len(data)} rows: {title}")
            return True

        except Exception as e:
            self.logger.error(f"Error creating table {title}: {e}")
            return False

    def safe_get_nested_data(self, data: Dict, path: str, default=None) -> Any:
        """
        Safely get nested data from dictionary using dot notation

        Args:
            data: Dictionary to search
            path: Dot-separated path (e.g., 'section.subsection.key')
            default: Default value if path not found

        Returns:
            Value at path or default
        """
        try:
            keys = path.split('.')
            current = data

            for key in keys:
                if isinstance(current, dict) and key in current:
                    current = current[key]
                else:
                    return default

            return current

        except Exception as e:
            self.logger.debug(f"Error getting nested data for path {path}: {e}")
            return default

    def process_technical_architecture(self, tech_arch_data: Dict):
        """Process technical architecture section with error handling"""
        try:
            self.doc.add_heading('2. Technical Architecture', level=1)

            # 2.1 Technology Stack
            project_access = self.safe_get_nested_data(tech_arch_data, 'project_access', {})
            tech_stack = self.safe_get_nested_data(project_access, '2.02 Technology Stack')

            if tech_stack:
                self.doc.add_heading('2.1 Technology Stack', level=2)
                tech_list = tech_stack.split(',') if ',' in tech_stack else [tech_stack]
                for tech in tech_list:
                    p = self.doc.add_paragraph(f'• {tech.strip()}')
                    p.paragraph_format.left_indent = Inches(0.25)

            # 2.2 Project Access
            pa_data = self.safe_get_nested_data(project_access, '2.04 Project Access', [])
            if pa_data:
                self.doc.add_heading('2.2 Project Access', level=2)
                headers = ['Group', 'Owners', 'Description', 'Environment', 'Environment Type', 'Region']
                self.create_professional_table(headers, pa_data)

            # 2.3 Key Components
            kc_data = self.safe_get_nested_data(tech_arch_data,
                                                'key_components.2.05 Key Components.- **2.5.1 Components**', [])
            if kc_data:
                self.doc.add_heading('2.3 Key Components', level=2)
                headers = ['Service', 'Dev', 'QA', 'Prod']
                self.create_professional_table(headers, kc_data)

            # 2.4 Service Accounts
            sa_data = self.safe_get_nested_data(tech_arch_data, 'service_accounts.2.09 Service Accounts', [])
            if sa_data:
                self.doc.add_heading('2.4 Service Accounts', level=2)
                headers = ['Service Account', 'Service', 'Environment', 'Permission', 'JSON Key']
                self.create_professional_table(headers, sa_data)

            # 2.5 Firewall Rules
            fw_data = self.safe_get_nested_data(tech_arch_data, 'firewall.2.12 Firewall Rules', [])
            if fw_data:
                self.doc.add_heading('2.5 Firewall Rules', level=2)
                headers = ['Network', 'Actions', 'Source IP', 'Destination', 'Port', 'Protocol', 'Direction']
                self.create_professional_table(headers, fw_data)

            # 2.6 IAM
            iam_data = self.safe_get_nested_data(tech_arch_data, 'iam.2.13 IAM', [])
            if iam_data:
                self.doc.add_heading('2.6 IAM', level=2)
                headers = ['Service', 'Principal', 'Environment', 'Permission']
                self.create_professional_table(headers, iam_data)

            self.logger.debug("Technical architecture section processed")

        except Exception as e:
            self.logger.error(f"Error processing technical architecture: {e}")

    def process_devops_operations(self, devops_data: Dict):
        """Process DevOps and Operations section with error handling"""
        try:
            self.doc.add_heading('3. Development and Operations', level=1)

            # 3.1 & 3.2 Development Environment and Deployment Process
            devops_section = self.safe_get_nested_data(devops_data, 'devops.3. Development and Operations', {})

            if devops_section:
                dev_env = devops_section.get('3.1 Development Environment', '')
                if dev_env:
                    self.doc.add_heading('3.1 Development Environment', level=2)
                    self.doc.add_paragraph(dev_env)

                deploy_process = devops_section.get('3.2 Deployment Process', '')
                if deploy_process:
                    self.doc.add_heading('3.2 Deployment Process', level=2)
                    self.doc.add_paragraph(deploy_process)

            # 3.3 Operational Support Model
            osm_data = self.safe_get_nested_data(devops_data, 'ops.3.3 Operational Support Model', [])
            if osm_data:
                self.doc.add_heading('3.3 Operational Support Model', level=2)
                headers = ['Component', 'DEV', 'QA', 'UAT', 'PROD']
                self.create_professional_table(headers, osm_data)

            # 3.4 Service Level Management
            service_data = self.safe_get_nested_data(devops_data, 'service', {})
            if service_data:
                self.doc.add_heading('3.4 Service Level Management', level=2)

                sli = service_data.get('3.4.1 Service Level Indicators (SLI)', '')
                if sli:
                    self.doc.add_heading('3.4.1 Service Level Indicators (SLI)', level=3)
                    self.doc.add_paragraph(sli)

                slo = service_data.get('3.4.2 Service Level Objective (SLO)', '')
                if slo:
                    self.doc.add_heading('3.4.2 Service Level Objective (SLO)', level=3)
                    self.doc.add_paragraph(slo)

            # 3.5 Governance and Cost Management
            governance_data = self.safe_get_nested_data(devops_data, 'governance', {})
            if governance_data:
                self.doc.add_heading('3.5 Governance and Cost Management', level=2)

                rpo = governance_data.get('3.4.3 Recovery Point Objective (RPO)', '')
                if rpo:
                    self.doc.add_heading('3.5.1 Recovery Point Objective (RPO)', level=3)
                    self.doc.add_paragraph(rpo)

                rto = governance_data.get('3.4.4 Recovery Time Objective (RTO)', '')
                if rto:
                    self.doc.add_heading('3.5.2 Recovery Time Objective (RTO)', level=3)
                    self.doc.add_paragraph(rto)

                gov = governance_data.get('3.6 Governance', '')
                if gov:
                    self.doc.add_heading('3.5.3 Governance', level=3)
                    self.doc.add_paragraph(gov)

                # Cost Management Table
                cost_data = governance_data.get('3.7 Cost Management / GCP Budget Approval', [])
                if cost_data:
                    self.doc.add_heading('3.5.4 Cost Management / GCP Budget Approval', level=3)
                    headers = ['Environment', 'Cost (in Dollars)']
                    self.create_professional_table(headers, cost_data)

            self.logger.debug("DevOps and operations section processed")

        except Exception as e:
            self.logger.error(f"Error processing DevOps operations: {e}")

    def process_risks_assumptions(self, risks_data: Dict):
        """Process Risks, Assumptions, and Decisions section with error handling"""
        try:
            self.doc.add_heading('4. Risks, Assumptions, and Decisions', level=1)

            # 4.1 Design Assumptions
            assumptions_path = 'Design_Assumptions.4. Risks, Assumptions, and Decisions.4.1 Design Assumptions'
            assumptions_data = self.safe_get_nested_data(risks_data, assumptions_path, [])
            if assumptions_data:
                self.doc.add_heading('4.1 Design Assumptions', level=2)
                headers = ['Assumption ID', 'Assumption', 'Validation Criteria']
                self.create_professional_table(headers, assumptions_data)

            # 4.2 Dependencies
            deps_data = self.safe_get_nested_data(risks_data, 'Dependencies.4.2 Dependencies', [])
            if deps_data:
                self.doc.add_heading('4.2 Dependencies', level=2)
                headers = ['Dependency', 'Provider', 'Reason', 'Date Required']
                self.create_professional_table(headers, deps_data)

            # 4.3 Design Risks
            risks_list = self.safe_get_nested_data(risks_data, 'Design_Risks.4.3 Design Risks', [])
            if risks_list:
                self.doc.add_heading('4.3 Design Risks', level=2)
                headers = ['ID', 'Category', 'Likelihood', 'Risks Impacts and Actions', 'Date Accepted/Removed']
                self.create_professional_table(headers, risks_list)

            # 4.4 Key Decisions
            decisions_data = self.safe_get_nested_data(risks_data, 'Key_Decisions.4.4 Key Decisions', [])
            if decisions_data:
                self.doc.add_heading('4.4 Key Decisions', level=2)
                headers = ['ID', 'Category', 'Decision Description', 'Date Raised', 'Date Accepted/Removed',
                           'Likelihood']
                self.create_professional_table(headers, decisions_data)

            self.logger.debug("Risks and assumptions section processed")

        except Exception as e:
            self.logger.error(f"Error processing risks and assumptions: {e}")

    def process_appendix(self, appendix_data: Dict):
        """Process Appendix section with error handling"""
        try:
            self.doc.add_heading('5. Appendix', level=1)

            # 5.1 Abbreviations and Glossary
            self.doc.add_heading('5.1 Abbreviations and Glossary', level=2)

            # 5.1.1 Abbreviations
            abbrev_path = 'Abbreviations.5. Appendix.5.1 Abbreviations and Glossary.5.1.1 Abbreviations'
            abbrev_data = self.safe_get_nested_data(appendix_data, abbrev_path, [])
            if abbrev_data:
                self.doc.add_heading('5.1.1 Abbreviations', level=3)
                headers = ['Abbreviation/Acronym', 'Description']
                self.create_professional_table(headers, abbrev_data)

            # 5.1.2 Glossary
            glossary_data = self.safe_get_nested_data(appendix_data, 'Glossary.5.1.2 Glossary', [])
            if glossary_data:
                self.doc.add_heading('5.1.2 Glossary', level=3)
                headers = ['Term', 'Definition']
                self.create_professional_table(headers, glossary_data)

            # 5.1.3 References
            refs_data = self.safe_get_nested_data(appendix_data, 'References.5.1.3 References', {})
            if refs_data:
                self.doc.add_heading('5.1.3 References', level=3)

                project_refs = refs_data.get('5.1.3.1 Required Project References', [])
                if project_refs:
                    self.doc.add_heading('5.1.3.1 Required Project References', level=4)
                    for ref in project_refs:
                        p = self.doc.add_paragraph(f'• {ref}')
                        p.paragraph_format.left_indent = Inches(0.25)

                other_refs = refs_data.get('5.1.3.2 Standards and other References', [])
                if other_refs:
                    self.doc.add_heading('5.1.3.2 Standards and other References', level=4)
                    for ref in other_refs:
                        p = self.doc.add_paragraph(f'• {ref}')
                        p.paragraph_format.left_indent = Inches(0.25)

            # 5.2 Resources and Guidance
            resources_data = self.safe_get_nested_data(appendix_data, 'Resources_Guidance.5.2 Resources and Guidance',
                                                       [])
            if resources_data:
                self.doc.add_heading('5.2 Resources and Guidance', level=2)
                headers = ['Resource', 'Link']
                self.create_professional_table(headers, resources_data)

            # 5.3 Google APIs
            apis_data = self.safe_get_nested_data(appendix_data, 'Resources_Guidance.5.3 Google APIs', [])
            if apis_data:
                self.doc.add_heading('5.3 Google APIs', level=2)
                headers = ['Service', 'API']
                self.create_professional_table(headers, apis_data)

            self.logger.debug("Appendix section processed")

        except Exception as e:
            self.logger.error(f"Error processing appendix: {e}")

    def add_headers_footers(self):
        """Add professional headers and footers to the document"""
        try:
            section = self.doc.sections[0]

            # Header
            header = section.header
            header_paragraph = header.paragraphs[0]
            header_paragraph.text = "Application Design Document"
            header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Make header text bold
            for run in header_paragraph.runs:
                run.bold = True

            # Add border to header
            self.add_border_to_paragraph(header_paragraph)

            # Footer
            footer = section.footer

            # Add border to footer
            footer_paragraph = footer.add_paragraph()
            self.add_border_to_paragraph(footer_paragraph)

            # Footer content
            p1 = footer.add_paragraph("Draft\nInternal Use Only")
            p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p1.runs:
                run.font.size = Pt(9)

            p2 = footer.add_paragraph("Architecture Design Document ©2024 Lumen Technologies. All Rights Reserved.")
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p2.runs:
                run.font.size = Pt(9)

            # Page numbers
            p3 = footer.add_paragraph("Page ")
            p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in p3.runs:
                run.font.size = Pt(9)

            self.logger.debug("Headers and footers added")

        except Exception as e:
            self.logger.error(f"Error adding headers/footers: {e}")

    def generate_add_document(self, add_data: Dict) -> List[str]:
        """
        Main method to generate the complete ADD document from agent output

        Args:
            add_data: Dictionary containing ADD data for one or more applications

        Returns:
            List of generated file names
        """
        generated_files = []

        try:
            self.logger.info(f"Starting document generation for {len(add_data)} application(s)")

            # Process each application in the data
            for app_name, app_data in add_data.items():
                self.logger.info(f"Generating document for application: {app_name}")

                try:
                    # Create title page
                    self.create_title_page(app_name, app_name)

                    # Add table of contents
                    self.add_table_of_contents()

                    # Add application overview
                    self.doc.add_heading('1. Application Overview', level=1)
                    overview_text = f'This document describes the comprehensive architecture design for the {app_name} application. '
                    overview_text += f'It covers technical architecture, development and operations procedures, risk assessments, and supporting documentation.'
                    self.doc.add_paragraph(overview_text)

                    # Add purpose and scope
                    self.doc.add_heading('1.1 Purpose and Scope', level=2)
                    purpose_text = f'The purpose of this document is to provide a detailed architectural blueprint for {app_name}, '
                    purpose_text += 'enabling stakeholders to understand the system design, deployment strategy, and operational requirements.'
                    self.doc.add_paragraph(purpose_text)

                    # Process each major section
                    if 'technical_architecture' in app_data:
                        self.process_technical_architecture(app_data['technical_architecture'])

                    if 'devops_and_operations' in app_data:
                        self.process_devops_operations(app_data['devops_and_operations'])

                    if 'risks_and_assumptions' in app_data:
                        self.process_risks_assumptions(app_data['risks_and_assumptions'])

                    if 'appendix' in app_data:
                        self.process_appendix(app_data['appendix'])

                    # Add headers and footers
                    self.add_headers_footers()

                    # Generate filename with timestamp
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    filename = f"{app_name}_Architecture_Design_Document_v1.4_{timestamp}.docx"

                    # Save the document
                    self.doc.save(filename)
                    generated_files.append(filename)

                    self.logger.info(f"Document generated successfully: {filename}")

                    # If multiple applications, create separate documents
                    if len(add_data) > 1:
                        self.doc = Document()  # Reset for next application
                        self.setup_styles()

                except Exception as e:
                    self.logger.error(f"Error generating document for {app_name}: {e}")
                    continue

            self.logger.info(f"Document generation completed. Generated {len(generated_files)} file(s)")
            return generated_files

        except Exception as e:
            self.logger.error(f"Critical error in document generation: {e}")
            return generated_files

    def upload_to_gcs(self, filename: str, bucket_name: str = None, folder_prefix: str = "add_documents") -> Optional[
        str]:
        """
        Upload file to Google Cloud Storage

        Args:
            filename: Local filename to upload
            bucket_name: GCS bucket name
            folder_prefix: Folder prefix in the bucket

        Returns:
            GCS path if successful, None if failed
        """
        try:
            from google.cloud import storage

            # Validate inputs
            if not bucket_name:
                bucket_name = os.environ.get('GCS_BUCKET_NAME')
                if not bucket_name:
                    self.logger.error("No GCS bucket name provided")
                    return None

            if not Path(filename).exists():
                self.logger.error(f"File not found: {filename}")
                return None

            # Initialize GCS client
            client = storage.Client()

            # Get bucket
            try:
                bucket = client.bucket(bucket_name)
                # Test bucket access
                bucket.exists()
            except Exception as e:
                self.logger.error(f"Cannot access GCS bucket '{bucket_name}': {e}")
                return None

            # Create blob name with timestamp and folder structure
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            blob_name = f"{folder_prefix}/{timestamp}/{filename}"
            blob = bucket.blob(blob_name)

            # Upload file with metadata
            blob.metadata = {
                'source': 'ADD_Generator',
                'generated_at': datetime.now().isoformat(),
                'content_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            }

            with open(filename, 'rb') as f:
                blob.upload_from_file(f,
                                      content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

            gcs_path = f"gs://{bucket_name}/{blob_name}"
            self.logger.info(f"File uploaded successfully to GCS: {gcs_path}")

            return gcs_path

        except ImportError:
            self.logger.error(
                "Google Cloud Storage library not available. Install with: pip install google-cloud-storage")
            return None
        except Exception as e:
            self.logger.error(f"Error uploading to GCS: {e}")
            return None

    def upload_multiple_to_gcs(self, filenames: List[str], bucket_name: str = None,
                               folder_prefix: str = "add_documents") -> List[str]:
        """
        Upload multiple files to Google Cloud Storage

        Args:
            filenames: List of local filenames to upload
            bucket_name: GCS bucket name
            folder_prefix: Folder prefix in the bucket

        Returns:
            List of successful GCS paths
        """
        uploaded_paths = []

        self.logger.info(f"Uploading {len(filenames)} file(s) to GCS")

        for filename in filenames:
            gcs_path = self.upload_to_gcs(filename, bucket_name, folder_prefix)
            if gcs_path:
                uploaded_paths.append(gcs_path)
            else:
                self.logger.warning(f"Failed to upload: {filename}")

        self.logger.info(f"Successfully uploaded {len(uploaded_paths)}/{len(filenames)} file(s)")
        return uploaded_paths

    def validate_add_data_structure(self, add_data: Dict) -> bool:
        """
        Validate ADD data structure before processing

        Args:
            add_data: ADD data dictionary

        Returns:
            bool: True if valid structure
        """
        try:
            if not isinstance(add_data, dict):
                self.logger.error("ADD data must be a dictionary")
                return False

            if not add_data:
                self.logger.error("ADD data is empty")
                return False

            required_sections = [
                'technical_architecture',
                'devops_and_operations',
                'risks_and_assumptions',
                'appendix'
            ]

            all_valid = True
            for app_name, app_data in add_data.items():
                if not isinstance(app_data, dict):
                    self.logger.error(f"Application data for '{app_name}' must be a dictionary")
                    all_valid = False
                    continue

                missing_sections = []
                for section in required_sections:
                    if section not in app_data:
                        missing_sections.append(section)

                if missing_sections:
                    self.logger.warning(f"Application '{app_name}' missing sections: {missing_sections}")
                    # Don't fail validation, just warn

                # Validate technical architecture structure
                if 'technical_architecture' in app_data:
                    tech_arch = app_data['technical_architecture']
                    if not isinstance(tech_arch, dict):
                        self.logger.error(f"Technical architecture for '{app_name}' must be a dictionary")
                        all_valid = False

            if all_valid:
                self.logger.debug("ADD data structure validation passed")
            else:
                self.logger.error("ADD data structure validation failed")

            return all_valid

        except Exception as e:
            self.logger.error(f"Error validating ADD data structure: {e}")
            return False

    def get_document_statistics(self, filename: str) -> Dict[str, Any]:
        """
        Get statistics about the generated document

        Args:
            filename: Path to the document file

        Returns:
            Dictionary with document statistics
        """
        try:
            if not Path(filename).exists():
                return {'error': 'File not found'}

            # File size
            file_size = Path(filename).stat().st_size

            # Load document to count elements
            doc = Document(filename)

            stats = {
                'filename': filename,
                'file_size_bytes': file_size,
                'file_size_mb': round(file_size / (1024 * 1024), 2),
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
                'sections': len(doc.sections),
                'created_at': datetime.now().isoformat(),
            }

            # Count table rows
            total_table_rows = sum(len(table.rows) for table in doc.tables)
            stats['total_table_rows'] = total_table_rows

            # Count non-empty paragraphs
            non_empty_paragraphs = sum(1 for p in doc.paragraphs if p.text.strip())
            stats['non_empty_paragraphs'] = non_empty_paragraphs

            self.logger.debug(f"Document statistics generated for {filename}")
            return stats

        except Exception as e:
            self.logger.error(f"Error getting document statistics: {e}")
            return {'error': str(e)}

    def cleanup_local_files(self, filenames: List[str], keep_latest: int = 0) -> None:
        """
        Clean up local files after processing

        Args:
            filenames: List of filenames to clean up
            keep_latest: Number of latest files to keep (0 = delete all)
        """
        try:
            if keep_latest > 0:
                # Sort by modification time and keep the latest ones
                file_info = []
                for f in filenames:
                    if Path(f).exists():
                        stat = Path(f).stat()
                        file_info.append((f, stat.st_mtime))

                file_info.sort(key=lambda x: x[1], reverse=True)
                files_to_delete = [f[0] for f in file_info[keep_latest:]]
            else:
                files_to_delete = [f for f in filenames if Path(f).exists()]

            deleted_count = 0
            for file_path in files_to_delete:
                try:
                    Path(file_path).unlink()
                    deleted_count += 1
                    self.logger.debug(f"Deleted local file: {file_path}")
                except Exception as e:
                    self.logger.warning(f"Could not delete {file_path}: {e}")

            self.logger.info(f"Cleaned up {deleted_count} local file(s)")

        except Exception as e:
            self.logger.error(f"Error during cleanup: {e}")

    def generate_document_with_upload(self,
                                      add_data: Dict,
                                      bucket_name: str = None,
                                      upload_to_gcs: bool = True,
                                      cleanup_local: bool = False,
                                      folder_prefix: str = "add_documents") -> Dict[str, Any]:
        """
        Complete workflow: Generate documents and optionally upload to GCS

        Args:
            add_data: ADD data dictionary
            bucket_name: GCS bucket name
            upload_to_gcs: Whether to upload to GCS
            cleanup_local: Whether to delete local files after upload
            folder_prefix: GCS folder prefix

        Returns:
            Dictionary with generation results and metadata
        """
        result = {
            'status': 'started',
            'generated_files': [],
            'gcs_paths': [],
            'statistics': [],
            'errors': [],
            'processing_time': 0
        }

        start_time = datetime.now()

        try:
            # Validate input data
            if not self.validate_add_data_structure(add_data):
                result['status'] = 'error'
                result['errors'].append('Invalid ADD data structure')
                return result

            # Generate documents
            self.logger.info("Starting document generation workflow")
            generated_files = self.generate_add_document(add_data)
            result['generated_files'] = generated_files

            if not generated_files:
                result['status'] = 'error'
                result['errors'].append('No documents were generated')
                return result

            # Generate statistics for each file
            for filename in generated_files:
                stats = self.get_document_statistics(filename)
                result['statistics'].append(stats)

            # Upload to GCS if requested
            if upload_to_gcs and bucket_name:
                self.logger.info("Starting GCS upload")
                gcs_paths = self.upload_multiple_to_gcs(generated_files, bucket_name, folder_prefix)
                result['gcs_paths'] = gcs_paths

                if len(gcs_paths) != len(generated_files):
                    result['errors'].append(f'Only {len(gcs_paths)}/{len(generated_files)} files uploaded successfully')

            # Cleanup local files if requested and upload was successful
            if cleanup_local and upload_to_gcs and result['gcs_paths']:
                self.cleanup_local_files(generated_files)
                result['local_files_cleaned'] = True

            # Calculate processing time
            end_time = datetime.now()
            result['processing_time'] = (end_time - start_time).total_seconds()

            # Set final status
            if result['errors']:
                result['status'] = 'completed_with_warnings'
            else:
                result['status'] = 'success'

            self.logger.info(f"Document generation workflow completed: {result['status']}")
            return result

        except Exception as e:
            result['status'] = 'error'
            result['errors'].append(str(e))
            result['processing_time'] = (datetime.now() - start_time).total_seconds()
            self.logger.error(f"Document generation workflow failed: {e}")
            return result


# Utility functions for external use
def create_word_generator(debug: bool = False) -> WordGenerator:
    """
    Factory function to create a WordGenerator instance

    Args:
        debug: Enable debug logging

    Returns:
        WordGenerator instance
    """
    return WordGenerator(debug=debug)


def generate_documents_from_add_data(add_data: Dict,
                                     config: Dict = None) -> Dict[str, Any]:
    """
    Convenience function to generate documents with configuration

    Args:
        add_data: ADD data dictionary
        config: Configuration dictionary with options

    Returns:
        Generation results
    """
    # Default configuration
    default_config = {
        'debug': False,
        'bucket_name': None,
        'upload_to_gcs': False,
        'cleanup_local': False,
        'folder_prefix': 'add_documents'
    }

    # Merge with provided config
    if config:
        default_config.update(config)

    # Create generator and process
    generator = WordGenerator(debug=default_config['debug'])

    return generator.generate_document_with_upload(
        add_data=add_data,
        bucket_name=default_config['bucket_name'],
        upload_to_gcs=default_config['upload_to_gcs'],
        cleanup_local=default_config['cleanup_local'],
        folder_prefix=default_config['folder_prefix']
    )


# Example usage and testing
if __name__ == "__main__":
    # Example usage of the WordGenerator
    sample_data = {
        "SAMPLE_APP": {
            "technical_architecture": {
                "project_access": {
                    "2.02 Technology Stack": "GKE",
                    "2.04 Project Access": [
                        {
                            "Group": "admin",
                            "Owners": "admin@example.com",
                            "Description": "Admin access",
                            "Environment": "dev, qa, prod",
                            "Environment type": "dedicated",
                            "Region": "us-central1"
                        }
                    ]
                }
            },
            "devops_and_operations": {},
            "risks_and_assumptions": {},
            "appendix": {}
        }
    }

    # Create generator
    generator = WordGenerator(debug=True)

    # Generate documents
    result = generator.generate_document_with_upload(
        add_data=sample_data,
        upload_to_gcs=False,  # Set to True if you want to test GCS upload
        cleanup_local=False
    )

    print("Generation Result:")
    print(json.dumps(result, indent=2, default=str))
